"""
ReportEngine - 研究报告生成
"""
import json
from typing import List, Dict, Optional
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from CodeEngine.ir.code_ir import CodeUnit, Category
from CodeEngine.selective_coding.agent import StoryLine
from utils.llm_client import llm_client


class ReportEngine:
    """报告生成引擎"""
    
    def __init__(self, output_dir: str = "./outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_open_coding_report(self, codes: List[CodeUnit], 
                                    source_info: str) -> str:
        """生成开放编码阶段报告"""
        
        report = f"""# 开放编码阶段报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}
**数据来源**: {source_info}

## 1. 编码概况

- 总编码数: {len(codes)}
- 编码来源: {len(set(c.text_segment.source for c in codes))} 个数据源

## 2. 初始代码列表

"""
        
        for i, code in enumerate(codes, 1):
            report += f"""### {i}. {code.code_label}
- **定义**: {code.code_definition}
- **置信度**: {code.confidence:.2f}
- **来源**: {code.created_by.value}
- **文本示例**: {code.text_segment.content[:100]}...

"""
        
        return report
    
    def generate_axial_coding_report(self, categories: List[Category],
                                     coding_matrix: Dict) -> str:
        """生成轴心编码阶段报告"""
        
        report = f"""# 轴心编码阶段报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}

## 1. 类属结构

共识别 {len(categories)} 个类属：

"""
        
        for cat in categories:
            report += f"""### {cat.name}
- **定义**: {cat.definition}
- **包含代码**: {len(cat.codes)} 个
- **属性**: {', '.join(cat.properties) if cat.properties else '无'}
- **维度**: {json.dumps(cat.dimensions, ensure_ascii=False) if cat.dimensions else '无'}

"""
        
        report += "\n## 2. 编码矩阵\n\n"
        report += "| 类属 | 包含代码 | 属性 |\n"
        report += "|------|----------|------|\n"
        
        for name, data in coding_matrix.items():
            codes_str = ', '.join([c['label'] for c in data['codes'][:5]])
            props_str = ', '.join(data['properties'][:3])
            report += f"| {name} | {codes_str}... | {props_str} |\n"
        
        return report
    
    def generate_selective_coding_report(self, core_category: Category,
                                         storyline: StoryLine,
                                         propositions: List[str]) -> str:
        """生成选择性编码阶段报告"""
        
        report = f"""# 选择性编码阶段报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}

## 1. 核心类属

**{core_category.name}**

{core_category.definition}

## 2. 理论故事线

### {storyline.title}

{storyline.narrative}

## 3. 关键类属

{', '.join(storyline.key_categories)}

## 4. 理论命题

"""
        
        for i, prop in enumerate(propositions, 1):
            report += f"{i}. {prop}\n"
        
        return report
    
    def generate_full_report(self, title: str, research_question: str,
                            codes: List[CodeUnit],
                            categories: List[Category],
                            core_category: Category,
                            storyline: StoryLine,
                            propositions: List[str]) -> str:
        """生成完整研究报告"""
        
        report = f"""# {title}

**研究问题**: {research_question}

**报告生成时间**: {datetime.now().strftime('%Y-%m-%d')}

---

## 摘要

本研究采用扎根理论方法，通过对数据的系统编码分析，构建了关于"{core_category.name}"的理论框架。

## 1. 研究方法

### 1.1 数据收集

### 1.2 分析方法

采用 Strauss & Corbin 的扎根理论方法论，进行三级编码：
- 开放编码（Open Coding）
- 轴心编码（Axial Coding）
- 选择性编码（Selective Coding）

## 2. 研究发现

### 2.1 开放编码结果

共生成 {len(codes)} 个初始代码。

### 2.2 轴心编码结果

形成 {len(categories)} 个类属：

"""
        
        for cat in categories:
            report += f"- **{cat.name}**: {cat.definition}\n"
        
        report += f"""

### 2.3 选择性编码结果

#### 核心类属：{core_category.name}

{core_category.definition}

#### 理论故事线

**{storyline.title}**

{storyline.narrative}

## 3. 理论贡献

### 3.1 理论命题

"""
        
        for i, prop in enumerate(propositions, 1):
            report += f"{i}. {prop}\n"
        
        report += """

### 3.2 与现有文献的对话

## 4. 讨论与结论

## 5. 研究局限与未来方向

---

*本报告由 CoderResearch 自动生成*
"""
        
        return report
    
    def export_to_docx(self, markdown_content: str, output_path: str) -> str:
        """导出为 Word 文档"""
        doc = Document()
        
        # 简单解析 markdown 并转换为 docx
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('|'):
                # 跳过表格分隔线
                if '---' not in line:
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        p = doc.add_paragraph()
                        p.add_run(' | '.join(cells))
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path
    
    def save_report(self, content: str, filename: str):
        """保存报告文件"""
        output_path = self.output_dir / filename
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return str(output_path)
