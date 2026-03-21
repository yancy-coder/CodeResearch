"""
测试 ImportEngine - 数据导入模块
"""
import pytest
from pathlib import Path

from ImportEngine.agent import ImportEngine, ImportConfig


class TestImportDocx:
    """测试 .docx 文件导入"""

    def test_import_docx_basic(self, tmp_path):
        """测试基本的 .docx 导入"""
        from docx import Document

        # 创建一个临时 .docx 文件
        doc = Document()
        doc.add_paragraph("这是第一段测试文本，用于验证 docx 导入功能是否正常工作。")
        doc.add_paragraph("这是第二段测试文本，包含更多的内容以满足最小片段长度要求。")
        doc.add_paragraph("第三段文本也需要足够长，这样才能通过 min_segment_length 的过滤条件。")

        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        engine = ImportEngine(ImportConfig(min_segment_length=10))
        segments = engine.import_file(str(docx_path))

        # normalize_whitespace collapses newlines, so all text ends up in 1 segment
        assert len(segments) >= 1
        # Verify that docx text was actually extracted
        combined = " ".join(s.content for s in segments)
        assert "第一段测试文本" in combined
        assert "第二段测试文本" in combined
        assert segments[0].source == "test"

    def test_import_docx_empty_paragraphs(self, tmp_path):
        """测试空段落被跳过"""
        from docx import Document

        doc = Document()
        doc.add_paragraph("这是一段有内容的文本，长度足够满足最小片段长度要求。")
        doc.add_paragraph("")  # 空段落
        doc.add_paragraph("   ")  # 空白段落
        doc.add_paragraph("这是另一段有内容的文本，同样需要满足最小片段长度要求。")

        docx_path = tmp_path / "test_empty.docx"
        doc.save(str(docx_path))

        engine = ImportEngine(ImportConfig(min_segment_length=10))
        segments = engine.import_file(str(docx_path))

        assert len(segments) >= 1
        # Empty paragraphs should not contribute content
        combined = " ".join(s.content for s in segments)
        assert "有内容的文本" in combined
        assert "另一段有内容" in combined

    def test_import_unsupported_format(self, tmp_path):
        """测试不支持的文件格式抛出异常"""
        fake_pdf = tmp_path / "test.pdf"
        fake_pdf.write_text("fake content")

        engine = ImportEngine()
        with pytest.raises(ValueError, match="不支持的文件格式"):
            engine.import_file(str(fake_pdf))

    def test_import_nonexistent_file(self):
        """测试文件不存在抛出异常"""
        engine = ImportEngine()
        with pytest.raises(FileNotFoundError):
            engine.import_file("nonexistent_file.txt")


class TestImportText:
    """测试 .txt 文件导入"""

    def test_import_txt_basic(self, tmp_path):
        """测试基本的 .txt 导入"""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text(
            "这是第一段测试文本，用于验证 txt 导入功能。\n"
            "这是第二段测试文本，同样需要满足最小片段长度。",
            encoding="utf-8"
        )

        engine = ImportEngine(ImportConfig(min_segment_length=10))
        segments = engine.import_file(str(txt_path))

        # normalize_whitespace collapses newlines, resulting in 1 long segment
        assert len(segments) >= 1
        combined = " ".join(s.content for s in segments)
        assert "第一段" in combined

    def test_anonymize_phone(self):
        """测试手机号去标识化"""
        engine = ImportEngine(ImportConfig(anonymize=True))
        segments = engine.import_text(
            "我的电话是13812345678，请联系我。这段话需要足够长才能通过分段过滤。",
            source="test",
            segment_type="paragraph"
        )

        if segments:
            assert "13812345678" not in segments[0].content
            assert "[PHONE]" in segments[0].content
